#!/usr/bin/env python3
"""
deepresearch/scripts/convert.py

Convert a deepresearch Markdown output file to PDF and/or DOCX,
rendering Mermaid code blocks as crisp images via mmdc.

  PDF  — Mermaid rendered as ultra-high-res PNG (3600px wide, ~648 DPI on A4)
  DOCX — Mermaid rendered as high-res PNG (2400px wide)

Both formats produce sharp diagrams at any practical zoom level.

Usage:
  python convert.py <input.md> [--format pdf|docx|both] [--out-dir <dir>]

Requirements:
  npm: @mermaid-js/mermaid-cli   (npm install -g @mermaid-js/mermaid-cli)
  pip: python-docx reportlab

Examples:
  python convert.py output/2026-03-19-skills-research.md
  python convert.py output/report.md --format pdf
  python convert.py output/report.md --format both --out-dir exports/
"""

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


# ── CLI ────────────────────────────────────────────────────────────────────────

def parse_args():
    p = argparse.ArgumentParser(description=__doc__,
                                formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("input",  type=Path, help="Source .md file")
    p.add_argument("--format", choices=["pdf", "docx", "both"], default="both")
    p.add_argument("--out-dir", type=Path, default=None,
                   help="Output directory (default: same as input file)")
    return p.parse_args()


# ── Mermaid rendering ──────────────────────────────────────────────────────────

def find_mmdc() -> str:
    """Return the mmdc executable path, checking Windows npm locations."""
    cmd = shutil.which("mmdc")
    if cmd:
        return cmd
    import os
    for c in [
        Path(os.environ.get("APPDATA", "")) / "npm" / "mmdc.cmd",
        Path(os.environ.get("APPDATA", "")) / "npm" / "mmdc",
    ]:
        if c.exists():
            return str(c)
    raise FileNotFoundError(
        "mmdc not found. Install: npm install -g @mermaid-js/mermaid-cli"
    )


def render_png(code: str, idx: int, tmpdir: Path, mmdc: str,
               width: int, suffix: str) -> Path:
    """Render a Mermaid block to PNG at the given pixel width."""
    mmd = tmpdir / f"d{idx:02d}.mmd"
    png = tmpdir / f"d{idx:02d}_{suffix}.png"
    mmd.write_text(code, encoding="utf-8")
    result = subprocess.run(
        [mmdc, "-i", str(mmd), "-o", str(png),
         "-t", "default", "-b", "white", "--width", str(width)],
        capture_output=True, text=True
    )
    if result.returncode != 0 or not png.exists():
        raise RuntimeError(f"mmdc failed for diagram {idx}:\n{result.stderr.strip()}")
    print(f"  [{suffix}] d{idx:02d}.png  {width}px  ({png.stat().st_size:,} bytes)")
    return png


# ── Markdown parser ────────────────────────────────────────────────────────────

def strip_inline(s: str) -> str:
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    return s


def parse_blocks(src: str, tmpdir: Path, mmdc: str,
                 need_pdf: bool, need_docx: bool) -> list:
    """
    Parse Markdown into typed blocks.
    Mermaid blocks → ("mermaid", pdf_png_or_None, docx_png_or_None)
    """
    blocks = []
    lines  = src.splitlines()
    i, idx = 0, 0

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code = "\n".join(code_lines)
            if lang == "mermaid":
                idx += 1
                pdf_png  = render_png(code, idx, tmpdir, mmdc, 3600, "pdf")  if need_pdf  else None
                docx_png = render_png(code, idx, tmpdir, mmdc, 2400, "docx") if need_docx else None
                blocks.append(("mermaid", pdf_png, docx_png))
            else:
                blocks.append(("code", lang, code))

        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("> "):
            blocks.append(("quote", line[2:].strip()))
        elif line.startswith("---"):
            blocks.append(("hr",))
        elif line.startswith("| "):
            tbl = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl.append(lines[i])
                i += 1
            blocks.append(("table", tbl))
            continue
        elif re.match(r"^[-*] ", line):
            blocks.append(("li", line[2:].strip()))
        elif line.strip() == "":
            blocks.append(("blank",))
        else:
            blocks.append(("p", line))

        i += 1

    return blocks


# ── DOCX builder ───────────────────────────────────────────────────────────────

def build_docx(blocks: list, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def _code_para(text):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F3F3F3")
        pPr.append(shd)
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    def _inline(para, text):
        for k, part in enumerate(re.split(r"`([^`]+)`", text)):
            run = para.add_run(part)
            if k % 2 == 1:
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)

    for block in blocks:
        kind = block[0]
        if kind == "h1":
            doc.add_heading(strip_inline(block[1]), level=1)
        elif kind == "h2":
            doc.add_heading(strip_inline(block[1]), level=2)
        elif kind == "h3":
            doc.add_heading(strip_inline(block[1]), level=3)
        elif kind == "quote":
            p = doc.add_paragraph(style="Quote")
            _inline(p, block[1])
        elif kind == "hr":
            doc.add_paragraph("─" * 60)
        elif kind == "mermaid":
            png = block[2]  # DOCX PNG
            if png:
                doc.add_picture(str(png), width=Inches(5.5))
                doc.add_paragraph()
        elif kind == "code":
            _code_para(block[2])
        elif kind == "table":
            rows = [r for r in block[1] if not re.match(r"^\|[-| ]+\|$", r)]
            if not rows:
                continue
            parsed = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
            ncols  = max(len(r) for r in parsed)
            table  = doc.add_table(rows=len(parsed), cols=ncols)
            table.style = "Table Grid"
            for ri, row in enumerate(parsed):
                for ci, cell in enumerate(row):
                    if ci < ncols:
                        tc = table.cell(ri, ci)
                        tc.text = strip_inline(cell)
                        if ri == 0:
                            for run in tc.paragraphs[0].runs:
                                run.bold = True
            doc.add_paragraph()
        elif kind == "li":
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, block[1])
        elif kind == "p":
            p = doc.add_paragraph()
            _inline(p, block[1])

    doc.save(out_path)
    print(f"DOCX → {out_path}")


# ── PDF builder ────────────────────────────────────────────────────────────────

def build_pdf(blocks: list, out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        Table, TableStyle, Preformatted, Image,
    )

    W, _ = A4
    styles = getSampleStyleSheet()
    MAX_W  = W - 6 * cm

    T  = ParagraphStyle("T",  parent=styles["Title"],
         fontSize=18, spaceAfter=4,   textColor=colors.HexColor("#1a1a1a"))
    H2 = ParagraphStyle("H2", parent=styles["Heading2"],
         fontSize=13, spaceBefore=10, spaceAfter=3,
         textColor=colors.HexColor("#333333"))
    H3 = ParagraphStyle("H3", parent=styles["Heading3"],
         fontSize=11, spaceBefore=8,  spaceAfter=2,
         textColor=colors.HexColor("#555555"))
    B  = ParagraphStyle("B",  parent=styles["Normal"],
         fontSize=10, leading=14, spaceAfter=6)
    Q  = ParagraphStyle("Q",  parent=styles["Normal"],
         fontSize=9,  leading=12, leftIndent=16,
         textColor=colors.HexColor("#666666"), spaceAfter=6)
    C  = ParagraphStyle("C",  parent=styles["Code"],
         fontSize=7.5, leading=11, backColor=colors.HexColor("#F3F3F3"),
         leftIndent=8, rightIndent=8, spaceBefore=4, spaceAfter=4)
    LI = ParagraphStyle("LI", parent=styles["Normal"],
         fontSize=10, leading=14, leftIndent=16,
         firstLineIndent=-10, spaceAfter=3)

    def esc(s):
        s = s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        s = re.sub(r"`([^`]+)`",
                   r'<font name="Courier" size="8">\1</font>', s)
        s = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", s)
        return s

    story = []
    for block in blocks:
        kind = block[0]
        if kind == "h1":
            story.append(Paragraph(esc(block[1]), T))
        elif kind == "h2":
            story.append(Paragraph(esc(block[1]), H2))
        elif kind == "h3":
            story.append(Paragraph(esc(block[1]), H3))
        elif kind == "quote":
            story.append(Paragraph(esc(block[1]), Q))
        elif kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5,
                color=colors.HexColor("#cccccc"), spaceAfter=6, spaceBefore=6))
        elif kind == "mermaid":
            png = block[1]  # high-res PDF PNG
            if png:
                img = Image(str(png))
                scale = MAX_W / img.imageWidth
                img.drawWidth  = img.imageWidth  * scale
                img.drawHeight = img.imageHeight * scale
                story.append(img)
                story.append(Spacer(1, 8))
        elif kind == "code":
            story.append(Preformatted(block[2], C))
        elif kind == "table":
            rows = [r for r in block[1]
                    if not re.match(r"^\|[-| ]+\|$", r)]
            if not rows:
                continue
            parsed = [[c.strip() for c in r.strip("|").split("|")]
                      for r in rows]
            ncols = max(len(r) for r in parsed)
            data  = [([strip_inline(c) for c in (row + [""]*ncols)[:ncols]])
                     for row in parsed]
            col_w = MAX_W / ncols
            t = Table(data, colWidths=[col_w]*ncols)
            t.setStyle(TableStyle([
                ("BACKGROUND",    (0,0),(-1,0), colors.HexColor("#E8E8E8")),
                ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
                ("FONTSIZE",      (0,0),(-1,-1), 8),
                ("GRID",          (0,0),(-1,-1), 0.5, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),
                 [colors.white, colors.HexColor("#FAFAFA")]),
                ("VALIGN",        (0,0),(-1,-1), "TOP"),
                ("TOPPADDING",    (0,0),(-1,-1), 4),
                ("BOTTOMPADDING", (0,0),(-1,-1), 4),
                ("LEFTPADDING",   (0,0),(-1,-1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))
        elif kind == "li":
            story.append(Paragraph("• " + esc(block[1]), LI))
        elif kind == "p":
            story.append(Paragraph(esc(block[1]), B))
        elif kind == "blank":
            story.append(Spacer(1, 4))

    SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=3*cm, rightMargin=3*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    ).build(story)
    print(f"PDF  → {out_path}")


# ── main ───────────────────────────────────────────────────────────────────────

def main():
    args = parse_args()

    src = args.input.resolve()
    if not src.exists():
        sys.exit(f"Error: {src} not found")

    out_dir = args.out_dir.resolve() if args.out_dir else src.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    stem     = src.stem
    out_pdf  = out_dir / f"{stem}.pdf"
    out_docx = out_dir / f"{stem}.docx"

    need_pdf  = args.format in ("pdf",  "both")
    need_docx = args.format in ("docx", "both")

    mmdc   = find_mmdc()
    tmpdir = Path(tempfile.mkdtemp(prefix="deepresearch_"))

    try:
        print(f"Source : {src}")
        print(f"Format : {args.format}  (PDF=3600px PNG, DOCX=2400px PNG)")
        print(f"Out dir: {out_dir}\n")
        print("Rendering Mermaid diagrams...")

        blocks = parse_blocks(
            src.read_text(encoding="utf-8"), tmpdir, mmdc,
            need_pdf=need_pdf, need_docx=need_docx,
        )
        print(f"Parsed {len(blocks)} blocks\n")

        if need_docx:
            build_docx(blocks, out_docx)
        if need_pdf:
            build_pdf(blocks, out_pdf)

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    print("\nDone.")


if __name__ == "__main__":
    main()
